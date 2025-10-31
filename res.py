
"""
Paper Explainer Pro – FAST Streamlit App (Single File)

✅ Core Features:
- PDF upload (1 or 2 papers)
- Section detection
- TL;DR summary
- Section explanations
- Interactive RAG chat
- Glossary
- Key contributions
- Comparison mode (two papers)
- Study mode (Research Gap Identifier)
- Export Markdown & DOCX

Run:
  pip install -U streamlit pymupdf google-generativeai sentence-transformers scikit-learn faiss-cpu numpy pandas python-docx
  streamlit run app.py
"""


import  io, re, json
from dataclasses import dataclass
from typing import List, Dict, Tuple

import numpy as np
import streamlit as st
import fitz  # PyMuPDF

# ---------------- Gemini ----------------
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

try:
    import google.generativeai as genai
    if GEMINI_API_KEY:
        genai.configure(api_key=GEMINI_API_KEY)
        _GEMINI = True
    else:
        _GEMINI = False
except Exception:
    _GEMINI = False

# ------------- Embedding Model -------------
try:
    from sentence_transformers import SentenceTransformer
    from sklearn.metrics.pairwise import cosine_similarity
    import faiss
    _EMB = True
except Exception:
    _EMB = False

# ------------- DOCX Export -------------
try:
    from docx import Document
    _DOCX = True
except Exception:
    _DOCX = False

# ------------- Data Types -------------
@dataclass
class Chunk:
    text: str
    section: str
    page: int

SECTION_KEYS = [
    "abstract","introduction","background","related work","methodology","methods",
    "experiments","results","analysis","discussion","conclusion","limitations","future work","references"
]
SEC_PATS = [re.compile(rf"^\s*(?P<title>{name})\b.*$", re.I|re.M) for name in SECTION_KEYS]
LANGS = ["English","Hindi","Gujarati","Hinglish","Basic English"]

# ------------- PDF Parsing -------------
def load_pdf_bytes(b: bytes) -> Tuple[List[str], str]:
    doc = fitz.open(stream=b, filetype="pdf")
    pages = [p.get_text("text") for p in doc]
    return pages, "\n".join(pages)

def detect_sections(full: str) -> Dict[str, str]:
    marks = []
    for pat in SEC_PATS:
        for m in pat.finditer(full):
            marks.append((m.start(), m.group("title").lower()))
    if not marks:
        return {"introduction": full}
    marks = sorted(set(marks))
    marks = [(0, "_start")] + marks + [(len(full), "_end")]
    out = {}
    for i in range(1, len(marks)-1):
        s_idx, title = marks[i]
        e_idx, _ = marks[i+1]
        if title in ("_start","_end"): continue
        out[title] = full[s_idx:e_idx].strip()
    return out

def chunk_text(sec: str, text: str, max_chars: int = 1500) -> List[Chunk]:
    paras = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks, buf, total = [], [], 0
    for p in paras:
        if total+len(p) > max_chars and buf:
            chunks.append(Chunk("\n\n".join(buf), sec, 1))
            buf, total = [p], len(p)
        else:
            buf.append(p); total += len(p)
    if buf: chunks.append(Chunk("\n\n".join(buf), sec, 1))
    return chunks

# ------------- Retriever -------------
class Retriever:
    def __init__(self):
        self.embedder = None; self.mat = None; self.texts = []
    def fit(self, chunks: List[Chunk]):
        self.texts = [c.text for c in chunks]
        if not _EMB: return
        self.embedder = SentenceTransformer("all-MiniLM-L6-v2")
        self.mat = self.embedder.encode(self.texts, normalize_embeddings=True).astype("float32")
        self.index = faiss.IndexFlatIP(self.mat.shape[1]); self.index.add(self.mat)
    def search(self, q: str, k: int = 5):
        if not _EMB:
            words = [w.lower() for w in re.findall(r"\w+", q)]
            scores = [(sum(t.lower().count(w) for w in words), i) for i,t in enumerate(self.texts)]
            return [i for _,i in sorted(scores, reverse=True)[:k]]
        qv = self.embedder.encode([q], normalize_embeddings=True).astype("float32")
        _, I = self.index.search(qv, k); return I[0].tolist()

# ------------- Gemini Call -------------
def gcall(prompt: str) -> str:
    if not _GEMINI:
        return "(No GEMINI_API_KEY set)\n\n" + prompt[:1000]
    model = genai.GenerativeModel("gemini-2.5-flash")
    try:
        resp = model.generate_content(prompt)
        return getattr(resp, "text", "") or "(No response)"
    except Exception as e:
        return f"(Gemini Error: {e})"

# ------------- Prompts -------------
def p_tldr(full: str, lang: str): 
    return f"Create a 5-7 bullet TL;DR in {lang}. Cover problem, method, data, results, contributions.\n\n{full[:16000]}"
def p_explain(section: str, content: str, level: str, lang: str): 
    return f"Explain '{section}' section at {level} level in {lang}. Include summary, key points, and relevance.\n\n{content[:8000]}"
def p_gloss(full: str, lang: str) -> str:
    return (
        f"Extract a glossary (Term | Definition) in markdown table, in {lang}. Limit to top 25 impactful terms.\n\n{full[:16000]}"
    )
def p_contrib(full: str, lang: str): 
    return f"List main contributions of this paper as concise bullets in {lang}.\n\n{full[:16000]}"
def p_gap(full: str, lang: str): 
    return f"Identify 5 research gaps in {lang} as JSON list of objects with fields: aspect, gap, suggestion.\n\n{full[:16000]}"
def p_compare(a: str, b: str, lang: str):
    return f"Compare the two research papers. Output in {lang} under headings: Similarities, Differences, and Recommendation.\n\nPaper A:\n{a[:15000]}\n\nPaper B:\n{b[:15000]}"
def p_qa(ctx: List[Tuple[str,str]], q: str, lang: str):
    context = "\n\n".join([f"[{s}] {t[:1000]}" for s,t in ctx])
    return f"Answer only from the context below in {lang}.\n\n{context}\n\nQuestion: {q}"

# ------------- Exports -------------
def build_md(title: str, tldr: str, sections: Dict[str,str], gloss: str, contrib: str) -> str:
    lines = [f"# {title}", "", "## TL;DR", tldr, "", "## Explanations"]
    for s, txt in sections.items():
        lines += [f"### {s.title()}", txt, ""]
    lines += ["## Key Contributions", contrib, "", "## Glossary", gloss]
    return "\n".join(lines)
def to_docx(md: str) -> bytes:
    if not _DOCX: return md.encode()
    doc = Document()
    for line in md.splitlines():
        if line.startswith("# "): doc.add_heading(line[2:], 1)
        elif line.startswith("## "): doc.add_heading(line[3:], 2)
        elif line.startswith("### "): doc.add_heading(line[4:], 3)
        else: doc.add_paragraph(line)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()

# ---------------- UI ----------------
st.set_page_config(page_title="Paper Explainer Pro (Fast)", page_icon="⚡", layout="wide")
st.title("⚡ Paper Explainer Pro – Fast Edition")


st.title("🎯 Research Gap Identification")

st.markdown("""
This section helps you **identify gaps in your study area** from your research text, paper summary, or literature review.
""")

# Text input for user to paste research content
text_input = st.text_area("Paste your research summary or paper abstract here:", height=200)

# Button to trigger analysis
if st.button("🔍 Identify Research Gaps"):
    if len(text_input.strip()) == 0:
        st.warning("⚠️ Please enter some text to analyze.")
    else:
        with st.spinner("Analyzing text for potential research gaps..."):
            # Simple logic-based approach (you can replace this with LLM call later)
            gap_phrases = [
                "few studies", "limited research", "little attention", "not explored",
                "gap in knowledge", "unexplored area", "lack of", "scarce evidence",
                "not been investigated", "remains unclear"
            ]

            found_gaps = [p for p in gap_phrases if re.search(p, text_input.lower())]

            if found_gaps:
                st.success("✅ Potential Gaps Identified:")
                for i, g in enumerate(found_gaps, 1):
                    st.markdown(f"{i}. **{g}**")
                st.info("These indicate possible directions for deeper investigation.")
            else:
                st.info("No direct research gap indicators found. Try expanding your text or using more technical language.")

# Optional expansion (for study notes)
with st.expander("🧠 Tips to Identify Gaps Manually"):
    st.markdown("""
    - Look for **contradictory findings** or **inconsistencies** across papers.  
    - Identify **understudied variables, populations, or methods**.  
    - Check for **dated references** (older than 5 years) — may signal research gaps.  
    - Look for phrases like *"however"*, *"nevertheless"*, *"remains unclear"* in your literature review.  
    """)

st.markdown("---")
st.caption("💡 Tip: Use this with your literature review to automatically highlight underexplored areas.")


with st.sidebar:
    st.header("Settings")
    lang = st.selectbox("Language", LANGS, 0)
    level = st.radio("Explanation Level", ["Basic","Intermediate","Advanced"], 1)
    file_a = st.file_uploader("📄 Paper A", type=["pdf"])
    file_b = st.file_uploader("📘 Paper B (Optional)", type=["pdf"])

if "state" not in st.session_state:
    st.session_state.state = {}

left, right = st.columns([3,1])

with left:
    if file_a:
        pages, full = load_pdf_bytes(file_a.read())
        st.success(f"Loaded: {file_a.name} ({len(pages)} pages)")
        sections = detect_sections(full)
        st.session_state.state.update({"sections_a": sections, "full_a": full})

        chunks = [c for s,t in sections.items() for c in chunk_text(s,t)]
        retr = Retriever(); retr.fit(chunks)
        st.session_state.state.update({"chunks_a": chunks, "retriever_a": retr})

        sec_list = list(sections.keys())
        sel = st.selectbox("View Section", sec_list)
        st.text_area("Preview", sections[sel][:1000], height=150)

        with st.expander("📌 TL;DR Summary", True):
            tldr = gcall(p_tldr(full, lang))
            st.markdown(tldr)
            st.session_state.state["tldr"] = tldr

        with st.expander("📘 Explain Section", False):
            sel2 = st.selectbox("Select Section", sec_list)
            if st.button("Explain"):
                out = gcall(p_explain(sel2, sections[sel2], level, lang))
                st.markdown(out)
                st.session_state.state["explain"] = out

        with st.expander("📗 Glossary & Contributions", False):
            gloss = gcall(p_gloss(full, lang))
            st.markdown(gloss)
            contrib = gcall(p_contrib(full, lang))
            st.markdown("### 🏁 Key Contributions")
            st.markdown(contrib)

        with st.expander("🔍 Research Gap Identifier", False):
            raw = gcall(p_gap(full, lang))
            try:
                items = json.loads(re.findall(r"\[.*\]", raw, re.S)[0])
                for i,it in enumerate(items):
                    st.markdown(f"**{i+1}. {it.get('aspect','Aspect')}**")
                    st.markdown(f"Gap: {it.get('gap','N/A')}")
                    st.markdown(f"Suggestion: {it.get('suggestion','N/A')}")
                    st.markdown("---")
            except: st.info("Could not parse research gaps.")

        # Export
        st.markdown("---")
        explained = {k: gcall(p_explain(k,v,level,lang))[:4000] for k,v in list(sections.items())[:3]}
        md = build_md(file_a.name, tldr, explained, gloss, contrib)
        st.download_button("⬇️ Download Notes (MD)", md.encode(), "notes.md")
        st.download_button("⬇️ Download Notes (DOCX)", to_docx(md), "notes.docx")

    else:
        st.info("Upload Paper A to begin.")

with right:
    st.subheader("💬 Chat with Paper")
    if file_a and "chunks_a" in st.session_state.state:
        if "chat" not in st.session_state: st.session_state.chat = []
        for m in st.session_state.chat:
            with st.chat_message(m["role"]): st.markdown(m["content"])
        q = st.chat_input("Ask about the paper…")
        if q:
            st.session_state.chat.append({"role":"user","content":q})
            retr = st.session_state.state["retriever_a"]
            chunks = st.session_state.state["chunks_a"]
            idxs = retr.search(q, 5)
            ctx = [(chunks[i].section, chunks[i].text) for i in idxs]
            ans = gcall(p_qa(ctx, q, lang))
            st.session_state.chat.append({"role":"assistant","content":ans})
            with st.chat_message("assistant"): st.markdown(ans)
    else:
        st.caption("Upload Paper A to enable chat.")

    st.subheader("🔁 Compare Papers")
    if file_a and file_b:
        _, fullB = load_pdf_bytes(file_b.read())
        comp = gcall(p_compare(st.session_state.state["full_a"], fullB, lang))
        st.markdown(comp)
